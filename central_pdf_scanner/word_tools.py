from __future__ import annotations

import shutil
import subprocess
import tempfile
import re
import contextlib
import logging
from io import BytesIO
from zipfile import ZIP_DEFLATED, ZipFile
from pathlib import Path

import fitz
from docx import Document
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree


VML_NS = "urn:schemas-microsoft-com:vml"
OFFICE_NS = "urn:schemas-microsoft-com:office:office"
WORD10_NS = "urn:schemas-microsoft-com:office:word"


class WordToolError(RuntimeError):
    pass


def pdf_to_word(input_pdf: str | Path, output_docx: str | Path, mode: str = "editable") -> Path:
    """Converte PDF para DOCX editavel, com alta fidelidade visual."""
    source = Path(input_pdf)
    if not source.is_file():
        raise WordToolError("PDF não encontrado.")
    target = Path(output_docx).with_suffix(".docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    if mode == "editable":
        if _convert_pdf_with_word(source, target):
            return target
        if _pdf_has_visible_text(source):
            try:
                return _pdf_to_word_portable_reflow(source, target)
            except Exception:
                # Conversor de seguranca para PDFs com estruturas incomuns.
                # Ele mantem o documento utilizavel mesmo quando o mecanismo de
                # reconstrucao em paragrafos nao consegue interpretar uma pagina.
                pass
            return _pdf_to_word_high_fidelity(source, target)
        return _pdf_to_word_editable(source, target)
    if mode == "visual":
        return _pdf_to_word_visual(source, target)
    raise WordToolError("Modo de conversão para Word inválido.")


def _configure_section(section, page: fitz.Page, *, margins: float) -> None:
    section.page_width = Inches(page.rect.width / 72.0)
    section.page_height = Inches(page.rect.height / 72.0)
    section.top_margin = Inches(margins)
    section.bottom_margin = Inches(margins)
    section.left_margin = Inches(margins)
    section.right_margin = Inches(margins)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)


def _remove_initial_paragraph(document: Document) -> None:
    if document.paragraphs:
        paragraph = document.paragraphs[0]
        paragraph._element.getparent().remove(paragraph._element)


def _convert_pdf_with_word(source: Path, target: Path) -> bool:
    """Usa o mecanismo PDF Reflow quando o Microsoft Word estiver instalado."""
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(source), ReadOnly=True, ConfirmConversions=False)
        document.SaveAs2(str(target), FileFormat=16)
        return target.is_file() and target.stat().st_size > 0
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def _pdf_has_visible_text(source: Path) -> bool:
    """Distingue texto normal da camada invisivel produzida por OCR."""
    pdf = fitz.open(source)
    try:
        for page in pdf:
            for span in page.get_texttrace():
                if int(span.get("type", 0)) != 3 and str(span.get("chars", "")):
                    return True
    finally:
        pdf.close()
    return False


def _compact_reflow_layout(source: Path, target: Path) -> None:
    """Evita pequenas quebras extras ao abrir o DOCX no LibreOffice.

    O PDF usa coordenadas exatas, enquanto Word e LibreOffice recalculam as
    metricas das fontes. Uma reducao discreta apenas nos espacamentos verticais
    compensa essa diferenca sem reduzir o tamanho das letras ou rasterizar texto.
    """
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with ZipFile(source) as input_archive, ZipFile(target, "w", ZIP_DEFLATED) as output_archive:
        for item in input_archive.infolist():
            data = input_archive.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                    changed = False
                    for spacing in root.iter(f"{{{word_namespace}}}spacing"):
                        for attribute in ("before", "after", "line"):
                            qualified = f"{{{word_namespace}}}{attribute}"
                            value = spacing.get(qualified)
                            if value and value.lstrip("-").isdigit():
                                compacted = str(max(0, round(int(value) * 0.94)))
                                if compacted != value:
                                    spacing.set(qualified, compacted)
                                    changed = True
                    if changed:
                        data = etree.tostring(
                            root,
                            xml_declaration=True,
                            encoding="UTF-8",
                            standalone=True,
                        )
                except (etree.XMLSyntaxError, ValueError):
                    pass
            output_archive.writestr(item, data)


def _pdf_to_word_portable_reflow(source: Path, target: Path) -> Path:
    """Reconstrui o PDF em paragrafos, tabelas e imagens editaveis.

    Este e o modo portatil de maior qualidade e funciona sem Microsoft Word,
    inclusive quando o arquivo final sera editado no LibreOffice.
    """
    from pdf2docx import Converter

    with tempfile.TemporaryDirectory(prefix="pdf_word_reflow_") as temporary:
        intermediate = Path(temporary) / "reflow.docx"
        converter = Converter(str(source))
        try:
            # O limite 1.0 evita descartar linhas legitimas em sumarios e
            # layouts densos. O modo sequencial tambem evita a corrida de
            # arquivos temporarios observada em documentos muito extensos.
            with open(Path(temporary) / "conversion.log", "w", encoding="utf-8") as quiet:
                logging.disable(logging.INFO)
                with contextlib.redirect_stdout(quiet), contextlib.redirect_stderr(quiet):
                    converter.convert(
                        str(intermediate),
                        line_overlap_threshold=1.0,
                        multi_processing=False,
                    )
        finally:
            logging.disable(logging.NOTSET)
            converter.close()
        if not intermediate.is_file() or intermediate.stat().st_size == 0:
            raise WordToolError("Não foi possível reconstruir o PDF no formato Word.")
        _compact_reflow_layout(intermediate, target)
    return target


def _clean_xml_text(value: str) -> str:
    return "".join(character for character in value if character in "\t\n\r" or ord(character) >= 32)


def _normalize_font_name(value: object) -> str:
    name = str(value or "Arial").split(",", 1)[0]
    name = re.sub(r"^[A-Z]{6}\+", "", name)
    name = re.sub(r"-(?:Bold|Italic|Oblique|Regular|Roman).*$", "", name, flags=re.IGNORECASE)
    if name.endswith("MT") and len(name) > 2:
        name = name[:-2]
    return name or "Arial"


def _textbox_run(span: dict) -> etree._Element:
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    original_font_name = str(span.get("font", "Arial"))
    font_name = _normalize_font_name(original_font_name)
    fonts = OxmlElement("w:rFonts")
    for attribute in ("ascii", "hAnsi", "eastAsia", "cs"):
        fonts.set(qn(f"w:{attribute}"), font_name)
    properties.append(fonts)
    size = max(5.0, min(72.0, float(span.get("size", 11))))
    size_tag = OxmlElement("w:sz")
    size_tag.set(qn("w:val"), str(max(10, round(size * 2))))
    properties.append(size_tag)
    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), str(max(10, round(size * 2))))
    properties.append(size_cs)
    flags = int(span.get("flags", 0))
    if flags & 16 or "bold" in original_font_name.lower():
        properties.append(OxmlElement("w:b"))
    if flags & 2 or "italic" in original_font_name.lower() or "oblique" in original_font_name.lower():
        properties.append(OxmlElement("w:i"))
    color_value = int(span.get("color", 0))
    color = OxmlElement("w:color")
    color.set(qn("w:val"), f"{color_value & 0xFFFFFF:06X}")
    properties.append(color)
    run.append(properties)
    text = OxmlElement("w:t")
    text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text.text = _clean_xml_text(str(span.get("text", "")))
    run.append(text)
    return run


def _append_textbox(paragraph, line: dict, page_width: float) -> None:
    spans = [span for span in line.get("spans", []) if str(span.get("text", ""))]
    if not spans:
        return
    x0, y0, x1, y1 = line["bbox"]
    largest_font = max(float(span.get("size", 11)) for span in spans)
    top = max(0.0, y0 - max(2.0, largest_font * 0.35))
    width = max(4.0, (x1 - x0) * 1.08 + 7.0)
    height = max(12.0, y1 - y0 + largest_font * 1.15 + 4.0)
    run = paragraph.add_run()
    pict = OxmlElement("w:pict")
    shape = etree.Element(
        f"{{{VML_NS}}}shape",
        nsmap={"v": VML_NS, "o": OFFICE_NS, "w10": WORD10_NS},
    )
    shape.set("type", "#_x0000_t202")
    shape.set("stroked", "f")
    shape.set("filled", "f")
    shape.set(
        "style",
        ";".join(
            (
                "position:absolute",
                f"margin-left:{x0:.2f}pt",
                f"margin-top:{top:.2f}pt",
                f"width:{min(width, page_width - x0):.2f}pt",
                f"height:{height:.2f}pt",
                "mso-position-horizontal-relative:page",
                "mso-position-vertical-relative:page",
                "mso-wrap-style:none",
                "z-index:251659264",
            )
        ),
    )
    wrap = etree.SubElement(shape, f"{{{WORD10_NS}}}wrap")
    wrap.set("type", "none")
    textbox = etree.SubElement(shape, f"{{{VML_NS}}}textbox")
    textbox.set("inset", "0,0,0,0")
    content = OxmlElement("w:txbxContent")
    text_paragraph = OxmlElement("w:p")
    paragraph_properties = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    paragraph_properties.append(spacing)
    text_paragraph.append(paragraph_properties)
    for span in spans:
        text_paragraph.append(_textbox_run(span))
    content.append(text_paragraph)
    textbox.append(content)
    pict.append(shape)
    run._r.append(pict)


def _remove_text_for_background(page: fitz.Page) -> None:
    for block in page.get_text("dict").get("blocks", []):
        if block.get("type") != 0:
            continue
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                if str(span.get("text", "")).strip():
                    rectangle = fitz.Rect(span["bbox"])
                    rectangle.x0 -= 0.4
                    rectangle.x1 += 0.4
                    page.add_redact_annot(rectangle, fill=None, cross_out=False)
    page.apply_redactions(images=0, graphics=0, text=0)


def _pdf_to_word_high_fidelity(source: Path, target: Path) -> Path:
    """Mantem o fundo da pagina e sobrepoe o texto em caixas editaveis.

    A estrategia funciona no LibreOffice e evita que o conteudo de uma pagina
    transborde para a seguinte, um problema comum em conversores por paragrafos.
    """
    pdf = fitz.open(source)
    background_pdf = fitz.open(source)
    document = Document()
    section = document.sections[0]
    try:
        first_page = pdf[0]
        _configure_section(section, first_page, margins=0)
        for page_index, page in enumerate(pdf):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            if page_index:
                paragraph.paragraph_format.page_break_before = True

            background_page = background_pdf[page_index]
            _remove_text_for_background(background_page)
            pixmap = background_page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
            background = BytesIO(pixmap.tobytes("png"))
            paragraph.add_run().add_picture(
                background,
                width=Pt(page.rect.width),
                height=Pt(max(1, page.rect.height - 1)),
            )
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    _append_textbox(paragraph, line, page.rect.width)
        document.save(target)
    finally:
        background_pdf.close()
        pdf.close()
    return target


def _pdf_to_word_editable(source: Path, target: Path) -> Path:
    """Recria texto, estilos basicos, imagens e espacamento em elementos editaveis."""
    pdf = fitz.open(source)
    document = Document()
    _remove_initial_paragraph(document)
    text_characters = 0
    try:
        for page_index, page in enumerate(pdf):
            if page_index > 0:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            section = document.sections[-1]
            margin = 0.28
            _configure_section(section, page, margins=margin)
            usable_width = page.rect.width - 2 * margin * 72
            previous_bottom = margin * 72
            blocks = sorted(page.get_text("dict").get("blocks", []), key=lambda item: (item["bbox"][1], item["bbox"][0]))
            for block in blocks:
                x0, y0, x1, y1 = block["bbox"]
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_before = Pt(max(0, min(36, y0 - previous_bottom)))
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.left_indent = Pt(max(0, x0 - margin * 72))
                # Reserva uma pequena folga para diferenças de métrica entre a fonte do PDF e a do Word.
                paragraph.paragraph_format.right_indent = Pt(max(0, page.rect.width - margin * 72 - x1 - 36))
                block_center = (x0 + x1) / 2
                if abs(block_center - page.rect.width / 2) < 18 and (x1 - x0) < usable_width * 0.88:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif x0 > page.rect.width * 0.55:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

                if block.get("type") == 0:
                    for line_index, line in enumerate(block.get("lines", [])):
                        if line_index:
                            paragraph.add_run().add_break()
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            text_characters += len(text.strip())
                            run = paragraph.add_run(text)
                            font_name = str(span.get("font", "Arial")).split(",", 1)[0]
                            run.font.name = font_name
                            run.font.size = Pt(max(5, min(72, float(span.get("size", 11)))))
                            flags = int(span.get("flags", 0))
                            run.bold = bool(flags & 16) or "bold" in font_name.lower()
                            run.italic = bool(flags & 2) or "italic" in font_name.lower()
                            color = int(span.get("color", 0))
                            run.font.color.rgb = RGBColor((color >> 16) & 255, (color >> 8) & 255, color & 255)
                elif block.get("type") == 1 and block.get("image"):
                    try:
                        width = min(max(0.25, (x1 - x0) / 72), usable_width / 72)
                        paragraph.add_run().add_picture(BytesIO(block["image"]), width=Inches(width))
                    except Exception:
                        paragraph._element.getparent().remove(paragraph._element)
                        continue
                else:
                    paragraph._element.getparent().remove(paragraph._element)
                    continue
                previous_bottom = max(previous_bottom, y1)
        if text_characters == 0:
            raise WordToolError(
                "Este PDF não possui texto reconhecido. Primeiro use 'PDF digitalizado para OCR' "
                "e depois converta o resultado para Word editável."
            )
        document.save(target)
    finally:
        pdf.close()
    return target


def _pdf_to_word_visual(source: Path, target: Path) -> Path:
    """Preserva visualmente cada pagina do PDF como imagem dentro do Word."""
    pdf = fitz.open(source)
    document = Document()
    with tempfile.TemporaryDirectory(prefix="pdf_word_visual_") as temp:
      try:
        for page_index, page in enumerate(pdf):
            if page_index > 0:
                document.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
            section = document.sections[-1]
            _configure_section(section, page, margins=0)
            image_path = Path(temp) / f"pagina_{page_index + 1:04d}.png"
            page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False).save(str(image_path))
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Inches(0)
            paragraph.paragraph_format.space_after = Inches(0)
            paragraph.add_run().add_picture(
                str(image_path), width=Inches(page.rect.width / 72.0), height=Inches(page.rect.height / 72.0)
            )
        document.save(target)
      finally:
        pdf.close()
    return target


def word_to_pdf(input_docx: str | Path, output_pdf: str | Path) -> Path:
    source = Path(input_docx).resolve()
    if not source.is_file():
        raise WordToolError("Documento Word não encontrado.")
    target = Path(output_pdf).with_suffix(".pdf").resolve()
    target.parent.mkdir(parents=True, exist_ok=True)

    if _convert_with_word(source, target):
        return target
    if _convert_with_libreoffice(source, target):
        return target
    raise WordToolError(
        "Não foi possível converter. Instale o Microsoft Word ou o LibreOffice. "
        "O programa tentará ambos automaticamente."
    )


def _convert_with_word(source: Path, target: Path) -> bool:
    try:
        import win32com.client  # type: ignore
    except ImportError:
        return False
    word = None
    document = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        document = word.Documents.Open(str(source), ReadOnly=True)
        document.ExportAsFixedFormat(str(target), 17)
        return target.is_file() and target.stat().st_size > 0
    except Exception:
        return False
    finally:
        if document is not None:
            try:
                document.Close(False)
            except Exception:
                pass
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass


def _convert_with_libreoffice(source: Path, target: Path) -> bool:
    office = shutil.which("soffice") or shutil.which("libreoffice")
    if not office:
        return False
    with tempfile.TemporaryDirectory(prefix="central_pdf_word_") as temp:
        command = [office, "--headless", "--convert-to", "pdf", "--outdir", temp, str(source)]
        result = subprocess.run(command, capture_output=True, text=True, timeout=180)
        generated = Path(temp) / f"{source.stem}.pdf"
        if result.returncode != 0 or not generated.is_file():
            return False
        shutil.copy2(generated, target)
    return target.is_file() and target.stat().st_size > 0
