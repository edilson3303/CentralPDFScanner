from __future__ import annotations

import json
import tempfile
import threading
import unittest
from unittest.mock import patch
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path

import fitz
from PIL import Image
from pypdf import PdfReader
from pypdf.constants import UserAccessPermissions

from central_pdf_scanner.pdf_tools import (
    PDFToolError,
    compose_scanned_pdf,
    crop_pdf,
    images_to_pdf,
    merge_pdfs,
    merge_pdf_pages,
    parse_page_spec,
    parse_split_intervals,
    pdf_to_jpg,
    protect_pdf,
    remove_pages,
    rotate_pages,
    save_preview_images_as_jpg,
    split_pdf,
    trim_vertical_pdf,
    unprotect_pdf,
)
from central_pdf_scanner.scanner import _detect_connection_type, filter_direct_scanners
from central_pdf_scanner.escl_scanner import (
    ESCLScannerError,
    detect_escl_sources,
    detect_escl_info,
    detect_escl_details,
    probe_escl_scanner,
    scan_escl_to_pdf,
    _scan_settings,
    validate_ip_settings,
)
from central_pdf_scanner.scanner import ScannerDevice, _serial_from_wia_properties, _sources_from_wia_capabilities
from central_pdf_scanner.word_tools import pdf_to_word
from central_pdf_scanner.ocr import find_tesseract, images_to_searchable_pdf
from central_pdf_scanner.scan_processing import is_blank_image, prepare_scanned_images
from central_pdf_scanner.diagnostics import build_scanner_diagnostic
from central_pdf_scanner.advanced_pdf import _find_libreoffice, compact_pdf, convert_pdf_to_pdfa, separate_pdf_batch
from central_pdf_scanner.progress import OperationCancelled
from central_pdf_scanner import app as application
from central_pdf_scanner.app import DEFAULT_SCAN_PROFILES, default_scan_basename
from central_pdf_scanner.security import (
    is_administrative_sid,
    launch_elevated_history,
    launch_elevated_settings,
)
from central_pdf_scanner.thumbnail_dialogs import adjusted_zoom, merge_window_size, selection_after_click
from central_pdf_scanner.word_tools import word_to_pdf
from central_pdf_scanner.scan_options import paper_size_escl_units, paper_size_pixels
from docx import Document


def make_pdf(path: Path, pages: int = 3) -> Path:
    document = fitz.open()
    for index in range(pages):
        page = document.new_page(width=595, height=842)
        page.insert_text((72, 100), f"Página de teste {index + 1}", fontsize=18)
    document.save(path)
    document.close()
    return path


class FakeESCLHandler(BaseHTTPRequestHandler):
    def log_message(self, _format, *_args) -> None:
        pass

    def do_GET(self) -> None:
        if self.path == "/eSCL/ScannerCapabilities":
            content = (
                b'<scan:ScannerCapabilities xmlns:scan="http://schemas.hp.com/imaging/escl/2011/05/03">'
                b'<scan:SerialNumber>LXM123456</scan:SerialNumber>'
                b'<scan:Platen><scan:PlatenInputCaps /></scan:Platen>'
                b'<scan:Adf><scan:AdfSimplexInputCaps /><scan:AdfDuplexInputCaps /></scan:Adf>'
                b'</scan:ScannerCapabilities>'
            )
            self.send_response(200)
            self.send_header("Content-Type", "application/xml")
            self.send_header("Content-Length", str(len(content)))
            self.end_headers()
            self.wfile.write(content)
            return
        if self.path == "/eSCL/ScanJobs/1/NextDocument":
            busy_responses = getattr(self.server, "document_busy_responses", 0)  # type: ignore[attr-defined]
            if busy_responses:
                self.server.document_busy_responses = busy_responses - 1  # type: ignore[attr-defined]
                self.send_error(503)
                return
            count = getattr(self.server, "documents_sent", 0)  # type: ignore[attr-defined]
            available = getattr(self.server, "documents_available", 1)  # type: ignore[attr-defined]
            if count >= available:
                self.send_error(404)
                return
            self.server.documents_sent = count + 1  # type: ignore[attr-defined]
            stream = BytesIO()
            Image.new("RGB", (320, 240), "white").save(stream, "JPEG")
            content = stream.getvalue()
            self.send_response(200)
            self.send_header("Content-Type", "image/jpeg")
            self.send_header("Content-Length", str(len(content)))
            self.end_headers()
            self.wfile.write(content)
            return
        self.send_error(404)

    def do_POST(self) -> None:
        if self.path != "/eSCL/ScanJobs":
            self.send_error(404)
            return
        conflicts = getattr(self.server, "job_conflicts", 0)  # type: ignore[attr-defined]
        if conflicts:
            self.server.job_conflicts = conflicts - 1  # type: ignore[attr-defined]
            self.send_error(409)
            return
        length = int(self.headers.get("Content-Length", "0"))
        self.server.scan_settings = self.rfile.read(length)  # type: ignore[attr-defined]
        self.server.documents_sent = 0  # type: ignore[attr-defined]
        self.send_response(201)
        self.send_header("Location", "/eSCL/ScanJobs/1")
        self.send_header("Content-Length", "0")
        self.end_headers()

    def do_DELETE(self) -> None:
        if self.path == "/eSCL/ScanJobs/1":
            self.server.job_released = True  # type: ignore[attr-defined]
            self.send_response(204)
            self.end_headers()
            return
        self.send_error(404)


class PDFToolsTests(unittest.TestCase):
    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name)
        self.source = make_pdf(self.root / "entrada.pdf")

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_parse_page_spec(self) -> None:
        self.assertEqual(parse_page_spec("1,3-4", 5), [0, 2, 3])
        with self.assertRaises(PDFToolError):
            parse_page_spec("0", 5)
        with self.assertRaises(PDFToolError):
            parse_page_spec("4-2", 5)

    def test_thumbnail_ctrl_and_shift_selection(self) -> None:
        selected, anchor = selection_after_click(set(), 2, None, ctrl=False, shift=False)
        self.assertEqual(selected, {2})
        selected, anchor = selection_after_click(selected, 4, anchor, ctrl=True, shift=False)
        self.assertEqual(selected, {2, 4})
        selected, anchor = selection_after_click(selected, 7, anchor, ctrl=False, shift=True)
        self.assertEqual(selected, {4, 5, 6, 7})

    def test_merge_window_is_sized_for_pages_and_screen(self) -> None:
        self.assertEqual(merge_window_size(1, 1920, 1080), (900, 520))
        self.assertEqual(merge_window_size(20, 1920, 1080), (955, 850))
        self.assertEqual(merge_window_size(20, 800, 600), (750, 530))

    def test_redaction_zoom_is_limited(self) -> None:
        self.assertEqual(adjusted_zoom(1.0, 1.25), 1.25)
        self.assertEqual(adjusted_zoom(4.0, 1.25), 4.0)
        self.assertEqual(adjusted_zoom(0.5, 0.8), 0.5)

    def test_parse_split_intervals(self) -> None:
        self.assertEqual(parse_split_intervals("1-2,3,4-5", 5), [[0, 1], [2], [3, 4]])
        with self.assertRaises(PDFToolError):
            parse_split_intervals("1-3,3-5", 5)

    def test_remove_pages(self) -> None:
        output = remove_pages(self.source, self.root / "removido.pdf", "2")
        self.assertEqual(len(PdfReader(str(output)).pages), 2)

    def test_merge_pdfs(self) -> None:
        second = make_pdf(self.root / "segundo.pdf", 2)
        output = merge_pdfs([self.source, second], self.root / "unido.pdf")
        self.assertEqual(len(PdfReader(str(output)).pages), 5)

    def test_rotate_pages(self) -> None:
        output = rotate_pages(self.source, self.root / "girado.pdf", 90, "2")
        pages = PdfReader(str(output)).pages
        self.assertEqual(pages[0].rotation, 0)
        self.assertEqual(pages[1].rotation, 90)

    def test_crop_pdf(self) -> None:
        output = crop_pdf(self.source, self.root / "cortado.pdf", 10, 10, 10, 10, "1")
        pages = PdfReader(str(output)).pages
        self.assertLess(float(pages[0].cropbox.width), 595)
        self.assertEqual(float(pages[1].cropbox.width), 595)

    def test_trim_vertical_pdf_in_centimeters(self) -> None:
        output = trim_vertical_pdf(self.source, self.root / "vertical.pdf", 1.0, 2.0, "1")
        pages = PdfReader(str(output)).pages
        self.assertAlmostEqual(float(pages[0].cropbox.height), 842 - (30 * 72 / 25.4), places=2)
        self.assertEqual(float(pages[1].cropbox.height), 842)

    def test_split_and_merge_selected_pages(self) -> None:
        outputs = split_pdf(self.source, self.root / "dividido", "1-2,3")
        self.assertEqual(len(outputs), 2)
        self.assertEqual(len(PdfReader(str(outputs[0])).pages), 2)
        joined = merge_pdf_pages([(self.source, 2), (self.source, 0)], self.root / "reordenado.pdf")
        reader = PdfReader(str(joined))
        self.assertIn("3", reader.pages[0].extract_text())
        self.assertIn("1", reader.pages[1].extract_text())

    def test_compose_scanned_pdf_reorders_and_rotates(self) -> None:
        output = compose_scanned_pdf(
            [(self.source, 2, 90), (self.source, 0, 0)],
            self.root / "digitalizacao_revisada.pdf",
        )
        reader = PdfReader(str(output))
        self.assertEqual(len(reader.pages), 2)
        self.assertIn("3", reader.pages[0].extract_text())
        self.assertEqual(reader.pages[0].rotation, 90)

    def test_save_preview_images_as_jpg_reorders_and_rotates(self) -> None:
        first = self.root / "primeira.jpg"
        second = self.root / "segunda.jpg"
        Image.new("RGB", (120, 60), "red").save(first)
        Image.new("RGB", (80, 140), "blue").save(second)
        outputs = save_preview_images_as_jpg(
            [(second, -1, 0), (first, -1, 90)], self.root / "jpg", "Scan_TESTE"
        )
        self.assertEqual([path.name for path in outputs], [
            "Scan_TESTE_pagina_001.jpg", "Scan_TESTE_pagina_002.jpg"
        ])
        with Image.open(outputs[1]) as rotated:
            self.assertEqual(rotated.size, (60, 120))

    def test_blank_page_detection_and_removal(self) -> None:
        blank = self.root / "branca.jpg"
        content = self.root / "conteudo.jpg"
        Image.new("RGB", (500, 700), "white").save(blank)
        image = Image.new("RGB", (500, 700), "white")
        for x in range(80, 420):
            for y in range(300, 320):
                image.putpixel((x, y), (0, 0, 0))
        image.save(content)
        with Image.open(blank) as opened:
            self.assertTrue(is_blank_image(opened))
        with Image.open(content) as opened:
            self.assertFalse(is_blank_image(opened))
        self.assertEqual(prepare_scanned_images([blank, content], remove_blank_pages=True), [content])

    def test_scan_processing_fast_path_preserves_jpeg_bytes(self) -> None:
        image = self.root / "pagina.jpg"
        Image.new("RGB", (1200, 1600), "white").save(image, "JPEG", quality=91)
        original = image.read_bytes()
        self.assertEqual(prepare_scanned_images([image]), [image])
        self.assertEqual(image.read_bytes(), original)

    def test_scanner_diagnostic_does_not_require_a_scanner(self) -> None:
        report = build_scanner_diagnostic("2.8.3", self.root)
        self.assertIn("Versão do software: 2.8.3", report)
        self.assertIn("SCANNERS INSTALADOS NO WINDOWS", report)
        self.assertIn("não contém imagens nem conteúdo", report)
        self.assertNotIn("Windows/Sistema:", report)
        self.assertNotIn("Arquitetura:", report)
        self.assertNotIn("Python interno:", report)
        self.assertNotIn("OCR Tesseract:", report)

    def test_scanner_diagnostic_tests_all_registered_network_scanners(self) -> None:
        scanners = [
            {"nome": "Protocolo", "ip": "192.168.1.20"},
            {"nome": "Arquivo", "ip": "192.168.1.21"},
        ]
        with patch(
            "central_pdf_scanner.diagnostics.detect_escl_details",
            side_effect=[
                ("Lexmark 1", "SERIE1", ("Platen",)),
                ("Lexmark 2", "SERIE2", ("Platen", "Feeder")),
            ],
        ) as detector:
            report = build_scanner_diagnostic("2.8.3", self.root, scanners)
        self.assertEqual(detector.call_count, 2)
        self.assertIn("1. Protocolo", report)
        self.assertIn("2. Arquivo", report)
        self.assertNotIn("192.168.1.20", report)
        self.assertNotIn("192.168.1.21", report)

    def test_compact_pdf_reports_before_and_after(self) -> None:
        image_path = self.root / "grande.jpg"
        image = Image.effect_noise((1400, 1800), 80).convert("RGB")
        image.save(image_path, "JPEG", quality=98)
        source = images_to_pdf([image_path], self.root / "grande.pdf")
        output, before, after = compact_pdf(source, self.root / "compactado.pdf", "Tamanho reduzido")
        self.assertTrue(output.is_file())
        self.assertEqual(before, source.stat().st_size)
        self.assertLessEqual(after, before)
        self.assertEqual(len(PdfReader(str(output)).pages), 1)

    def test_separate_batch_by_page_count(self) -> None:
        source = make_pdf(self.root / "lote.pdf", 5)
        outputs = separate_pdf_batch(source, self.root / "lotes", "Quantidade de páginas", 2)
        self.assertEqual([len(PdfReader(str(path)).pages) for path in outputs], [2, 2, 1])

    def test_separate_batch_by_blank_page(self) -> None:
        document = fitz.open()
        first = document.new_page()
        first.insert_text((72, 72), "Processo um")
        document.new_page()
        third = document.new_page()
        third.insert_text((72, 72), "Processo dois")
        source = self.root / "separadores.pdf"
        document.save(source)
        document.close()
        outputs = separate_pdf_batch(source, self.root / "separados", "Página em branco")
        self.assertEqual(len(outputs), 2)
        self.assertEqual([len(PdfReader(str(path)).pages) for path in outputs], [1, 1])

    def test_cancellable_operation_stops_before_writing(self) -> None:
        event = threading.Event()
        event.set()
        with self.assertRaises(OperationCancelled):
            compact_pdf(self.source, self.root / "cancelado.pdf", "Equilibrado", cancel_event=event)

    def test_configurable_scan_filename(self) -> None:
        name = default_scan_basename(
            "ABC 123", {"modelo_nome": "Scan_{serie}_{data}_{hora}_{setor}", "setor": "Protocolo"}
        )
        self.assertRegex(name, r"^Scan_ABC_123_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}_Protocolo$")

    def test_default_scan_filename_includes_sector(self) -> None:
        name = default_scan_basename("ABC 123", {"setor": "Arquivo"})
        self.assertRegex(name, r"^Scan_ABC_123_\d{4}-\d{2}-\d{2}_\d{2}-\d{2}-\d{2}_Arquivo$")

    def test_paper_size_conversions_and_escl_request(self) -> None:
        self.assertEqual(paper_size_pixels("A4 (210 × 297 mm)", 300), (2480, 3508))
        self.assertEqual(paper_size_escl_units("Carta (216 × 279 mm)"), (2550, 3300))
        settings = _scan_settings(300, "Cor", "Platen", "A4 (210 × 297 mm)")
        self.assertIn(b"<pwg:Width>2480</pwg:Width>", settings)
        self.assertIn(b"<pwg:Height>3508</pwg:Height>", settings)
        self.assertIn(b"ThreeHundredthsOfInches", settings)

    @unittest.skipUnless(_find_libreoffice(), "LibreOffice não instalado")
    def test_pdfa_conversion_identifies_pdfa_2b(self) -> None:
        output = convert_pdf_to_pdfa(self.source, self.root / "arquivo.pdf")
        document = fitz.open(output)
        try:
            metadata = document.get_xml_metadata()
        finally:
            document.close()
        self.assertIn("<pdfaid:part>2</pdfaid:part>", metadata)
        self.assertIn("<pdfaid:conformance>B</pdfaid:conformance>", metadata)

    def test_image_round_trip(self) -> None:
        image = self.root / "imagem.jpg"
        Image.new("RGB", (320, 240), "navy").save(image)
        pdf = images_to_pdf([image], self.root / "imagem.pdf")
        outputs = pdf_to_jpg(pdf, self.root / "jpg", 100)
        self.assertEqual(len(outputs), 1)
        self.assertTrue(outputs[0].is_file())

    def test_pdf_to_word(self) -> None:
        output = pdf_to_word(self.source, self.root / "saida.docx")
        self.assertTrue(output.is_file())
        self.assertGreater(output.stat().st_size, 0)
        document = Document(output)
        text = "".join(document.element.body.itertext())
        self.assertIn("Página de teste 1", text)
        self.assertGreaterEqual(len(document.element.xpath(".//w:t")), 3)

    def test_pdf_to_word_visual_mode(self) -> None:
        output = pdf_to_word(self.source, self.root / "visual.docx", "visual")
        document = Document(output)
        self.assertGreaterEqual(len(document.inline_shapes), 3)

    def test_protect_and_unprotect_pdf(self) -> None:
        protected = protect_pdf(self.source, self.root / "protegido.pdf", "Abrir123", "Dono123", True)
        encrypted_reader = PdfReader(str(protected))
        self.assertTrue(encrypted_reader.is_encrypted)
        self.assertEqual(int(encrypted_reader.decrypt("Abrir123")), 1)
        self.assertEqual(len(encrypted_reader.pages), 3)
        permissions = encrypted_reader.user_access_permissions
        self.assertFalse(bool(permissions & UserAccessPermissions.MODIFY))
        self.assertFalse(bool(permissions & UserAccessPermissions.ADD_OR_MODIFY))
        self.assertFalse(bool(permissions & UserAccessPermissions.EXTRACT))
        self.assertFalse(bool(permissions & UserAccessPermissions.EXTRACT_TEXT_AND_GRAPHICS))

        encrypted_reader = PdfReader(str(protected))
        self.assertEqual(int(encrypted_reader.decrypt("Dono123")), 2)

        unprotected = unprotect_pdf(protected, self.root / "sem_senha.pdf", "Dono123")
        open_reader = PdfReader(str(unprotected))
        self.assertFalse(open_reader.is_encrypted)
        self.assertEqual(len(open_reader.pages), 3)

    def test_unprotect_rejects_wrong_password(self) -> None:
        protected = protect_pdf(self.source, self.root / "protegido.pdf", "Senha123")
        with self.assertRaisesRegex(PDFToolError, "incorretas"):
            unprotect_pdf(protected, self.root / "falha.pdf", "errada")

    def test_unprotect_accepts_separate_opening_and_owner_passwords(self) -> None:
        protected = protect_pdf(
            self.source, self.root / "duas_senhas.pdf", "Abrir123", "Editar456", True
        )
        output = unprotect_pdf(
            protected, self.root / "desprotegido.pdf", "Abrir123", "Editar456"
        )
        self.assertFalse(PdfReader(str(output)).is_encrypted)

    def test_administrative_group_sids(self) -> None:
        self.assertTrue(is_administrative_sid("S-1-5-32-544"))
        self.assertTrue(is_administrative_sid("S-1-5-21-100-200-300-512"))
        self.assertTrue(is_administrative_sid("S-1-5-21-100-200-300-519"))
        self.assertFalse(is_administrative_sid("S-1-5-21-100-200-300-513"))

    def test_administrative_windows_use_separate_uac_routes(self) -> None:
        with patch("central_pdf_scanner.security.launch_elevated_mode", return_value=True) as launch:
            self.assertTrue(launch_elevated_history())
            launch.assert_called_once_with("--historico")
        with patch("central_pdf_scanner.security.launch_elevated_mode", return_value=True) as launch:
            self.assertTrue(launch_elevated_settings())
            launch.assert_called_once_with("--configuracoes")

    def test_default_scan_mode_is_color(self) -> None:
        self.assertEqual(DEFAULT_SCAN_PROFILES["Documento padrão"]["color"], "Cor")

    def test_opening_and_editing_protections_are_independent(self) -> None:
        opening_only = protect_pdf(self.source, self.root / "abertura.pdf", "Abrir123")
        reader = PdfReader(str(opening_only))
        self.assertEqual(int(reader.decrypt("")), 0)
        self.assertEqual(int(reader.decrypt("Abrir123")), 1)
        self.assertTrue(bool(reader.user_access_permissions & UserAccessPermissions.MODIFY))

        editing_only = protect_pdf(self.source, self.root / "edicao.pdf", "", "Dono123", True)
        reader = PdfReader(str(editing_only))
        self.assertEqual(int(reader.decrypt("")), 1)
        self.assertFalse(bool(reader.user_access_permissions & UserAccessPermissions.MODIFY))
        self.assertFalse(bool(reader.user_access_permissions & UserAccessPermissions.EXTRACT))

    def test_scanner_serial_number_detection(self) -> None:
        class Property:
            def __init__(self, name: str, value: str) -> None:
                self.Name = name
                self.Value = value

        properties = [Property("Nome", "Scanner"), Property("Número de série", "LXM987654")]
        self.assertEqual(_serial_from_wia_properties(properties), "LXM987654")
        device = ScannerDevice("id", "Lexmark", "Rede", ("Vidro",), "LXM987654")
        self.assertIn("Série: LXM987654", device.display_name)

    def test_scanner_connection_type_labels(self) -> None:
        self.assertEqual(_detect_connection_type(r"USBSCAN\\VID_1234"), "USB / conectado")
        self.assertEqual(_detect_connection_type("SWD DAFWSDProvider WSD Scanner"), "Rede")
        self.assertEqual(_detect_connection_type("Scanner virtual"), "Instalado no Windows")

    def test_scanner_usb_excludes_devices_identified_as_network(self) -> None:
        devices = [
            ScannerDevice("usb", "Scanner USB", "USB / conectado"),
            ScannerDevice("rede", "Scanner WSD", "Rede"),
            ScannerDevice("local", "Scanner local", "Instalado no Windows"),
        ]
        self.assertEqual(
            [device.device_id for device in filter_direct_scanners(devices)],
            ["usb", "local"],
        )

    def test_wia_scanner_source_detection(self) -> None:
        self.assertEqual(_sources_from_wia_capabilities(2), ("Vidro",))
        self.assertEqual(_sources_from_wia_capabilities(1), ("Alimentador superior - somente frente",))
        self.assertEqual(
            _sources_from_wia_capabilities(7),
            ("Vidro", "Alimentador superior - somente frente", "Alimentador superior - frente e verso"),
        )

    def test_ip_scanner_settings_validation(self) -> None:
        self.assertEqual(validate_ip_settings("192.168.1.50", 80, "HTTP"), ("192.168.1.50", 80, "http"))
        with self.assertRaises(ESCLScannerError):
            validate_ip_settings("impressora.local", 80, "http")
        with self.assertRaises(ESCLScannerError):
            validate_ip_settings("192.168.1.50", 70000, "http")
        with self.assertRaisesRegex(ESCLScannerError, "rede local"):
            validate_ip_settings("8.8.8.8", 80, "http")

    def test_network_scanners_come_only_from_administrative_settings(self) -> None:
        machine = self.root / "maquina"
        user = self.root / "usuario"
        machine.mkdir()
        user.mkdir()
        (user / "configuracao.json").write_text(
            json.dumps({"ultimo_ip_scanner": "192.168.1.99"}), encoding="utf-8"
        )
        original_machine = application.machine_settings_directory
        original_user = application.settings_directory
        try:
            application.machine_settings_directory = lambda: machine
            application.settings_directory = lambda: user
            self.assertEqual(application._load_network_scanners(), [])
            (machine / "configuracao.json").write_text(
                json.dumps({
                    "scanners_rede": [
                        {"nome": "Protocolo", "ip": "192.168.1.20"},
                        {"nome": "Recursos Humanos", "ip": "192.168.1.21"},
                    ]
                }),
                encoding="utf-8",
            )
            self.assertEqual(
                application._load_network_scanners(),
                [
                    {"nome": "Protocolo", "ip": "192.168.1.20"},
                    {"nome": "Recursos Humanos", "ip": "192.168.1.21"},
                ],
            )
        finally:
            application.machine_settings_directory = original_machine
            application.settings_directory = original_user

    def test_escl_rejects_invalid_job_location(self) -> None:
        from central_pdf_scanner.escl_scanner import _job_url

        with self.assertRaisesRegex(ESCLScannerError, "trabalho inválido"):
            _job_url("http://192.168.1.50:80", "http://example.com/outro/caminho")

    def test_escl_accepts_lexmark_job_location_variants(self) -> None:
        from central_pdf_scanner.escl_scanner import _job_url

        base = "http://192.168.1.50:80"
        self.assertEqual(_job_url(base, "/eSCL/ScanJobs"), base + "/eSCL/ScanJobs")
        self.assertEqual(
            _job_url(base, "http://lexmark.local/ESCL/SCANJOBS/123/"),
            base + "/ESCL/SCANJOBS/123",
        )
        self.assertEqual(
            _job_url(base, "http://lexmark.local/eSCL/ScanJob/173f53b8-3f7c-4fc6-9405-e95c04d359cd"),
            base + "/eSCL/ScanJob/173f53b8-3f7c-4fc6-9405-e95c04d359cd",
        )

    def test_escl_probe_and_scan_to_pdf(self) -> None:
        server = ThreadingHTTPServer(("127.0.0.1", 0), FakeESCLHandler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            port = server.server_address[1]
            server.documents_available = 3  # type: ignore[attr-defined]
            server.job_conflicts = 2  # type: ignore[attr-defined]
            self.assertIn("encontrado", probe_escl_scanner("127.0.0.1", port))
            self.assertEqual(detect_escl_sources("127.0.0.1", port), ("Platen", "Feeder", "FeederDuplex"))
            name, sources = detect_escl_info("127.0.0.1", port)
            self.assertEqual(name, "Scanner_127.0.0.1")
            self.assertEqual(sources, ("Platen", "Feeder", "FeederDuplex"))
            model, serial_number, detail_sources = detect_escl_details("127.0.0.1", port)
            self.assertEqual(model, "Scanner_127.0.0.1")
            self.assertEqual(serial_number, "LXM123456")
            self.assertEqual(detail_sources, sources)
            output = scan_escl_to_pdf(
                "127.0.0.1",
                port,
                "http",
                self.root / "scanner_ip.pdf",
                dpi=300,
                color_mode="Cor",
                input_source="FeederDuplex",
            )
            self.assertTrue(output.is_file())
            self.assertEqual(len(PdfReader(str(output)).pages), 3)
            settings = server.scan_settings  # type: ignore[attr-defined]
            self.assertIn(b"RGB24", settings)
            self.assertIn(b"300", settings)
            self.assertIn(b"Feeder", settings)
            self.assertIn(b"<scan:Duplex>true</scan:Duplex>", settings)
            self.assertTrue(server.job_released)  # type: ignore[attr-defined]
        finally:
            server.shutdown()
            server.server_close()

    def test_escl_platen_waits_for_temporary_busy_response(self) -> None:
        server = ThreadingHTTPServer(("127.0.0.1", 0), FakeESCLHandler)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            port = server.server_address[1]
            server.document_busy_responses = 2  # type: ignore[attr-defined]
            output = scan_escl_to_pdf(
                "127.0.0.1",
                port,
                "http",
                self.root / "vidro.pdf",
                input_source="Platen",
            )
            self.assertTrue(output.is_file())
            self.assertTrue(server.job_released)  # type: ignore[attr-defined]
        finally:
            server.shutdown()
            server.server_close()

    def test_ocr_searchable_pdf_when_available(self) -> None:
        if find_tesseract() is None:
            self.skipTest("Tesseract não disponível")
        image = self.root / "ocr.png"
        canvas = Image.new("RGB", (1200, 300), "white")
        from PIL import ImageDraw
        ImageDraw.Draw(canvas).text((80, 100), "TESTE OCR 12345", fill="black", font_size=64)
        canvas.save(image)
        output = images_to_searchable_pdf([image], self.root / "ocr.pdf", language="eng")
        text = "".join(page.extract_text() or "" for page in PdfReader(str(output)).pages)
        self.assertIn("12345", text)

    def test_word_to_pdf_when_office_available(self) -> None:
        document = Document()
        document.add_heading("Documento de teste", 0)
        document.add_paragraph("Conversão Word para PDF.")
        source = self.root / "word.docx"
        document.save(source)
        try:
            output = word_to_pdf(source, self.root / "word.pdf")
        except Exception as exc:
            self.skipTest(f"Office/LibreOffice não disponível: {exc}")
        self.assertEqual(len(PdfReader(str(output)).pages), 1)


if __name__ == "__main__":
    unittest.main()
